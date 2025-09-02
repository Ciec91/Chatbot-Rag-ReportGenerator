#!/usr/bin/env python
# coding: utf-8

import os
import shutil
from dotenv import load_dotenv
from langchain.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI
import gradio as gr
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import json

# ================= CONFIGURATION =================

# Base directory path
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

load_dotenv()
openai_api_key = os.getenv('OPENAI_API_KEY')
if not openai_api_key:
    raise ValueError("❌ OPENAI_API_KEY is not configured")

client = OpenAI(api_key=openai_api_key)
MODEL = 'gpt-4o'

pdf_dir = '../CyberAudit_Rag/cybersecurity_pdf_data'
persist_path = os.path.expanduser("~/chroma_index_rag")

os.makedirs(persist_path, exist_ok=True)

if not os.access(persist_path, os.W_OK):
    raise PermissionError(f"Cannot write to the directory {persist_path}")

all_docs = []
for filename in os.listdir(pdf_dir):
    if filename.endswith(".pdf"):
        path = os.path.join(pdf_dir, filename)
        try:
            loader = PyPDFLoader(path)
            docs = loader.load()
            all_docs.extend(docs)
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=10)
chunks = splitter.split_documents(all_docs)
print(f"📑 Chunks generated: {len(chunks)}")

embedding = OpenAIEmbeddings(openai_api_key=openai_api_key)
vectorstore = Chroma.from_documents(
    chunks, embedding, persist_directory=persist_path)
print("✅ Index created and saved.")

system_message = """
You are a highly specialized knowledge assistant in **cybersecurity compliance audits** for **BDO Singapore**.

Your role is to support BDO staff by providing accurate, document-based responses related to audit and regulatory requirements. Always respond in the **same language** used by the user.

You must base your responses **strictly** on the content of the provided document set, which includes:

- **Singapore cybersecurity and data protection regulations**, such as the Personal Data Protection Act (PDPA) and PDPC guidelines.
- **Monetary Authority of Singapore (MAS) requirements**, including the Technology Risk Management (TRM) Guidelines, MAS Notice 626, and other applicable circulars.
- **International standards and frameworks** relevant to audits, such as ISO/IEC 27001, NIST Cybersecurity Framework, and COBIT.
- **BDO Singapore internal policies**, covering cybersecurity, risk management, data governance, and audit procedures.
- **BDO Singapore’s operational manuals and procedures** related to incident response, access control, data retention, and other cybersecurity processes.

When answering:

- Be **precise, concise, and professional**.
- Use **bullet points** for clarity whenever possible.
- **Cite the source document explicitly**, including the document title, section name, and any relevant identifiers.
- **Never provide information not contained in the documents**. Do not infer or fabricate content.
- If the documents do not provide an answer, clearly state that **there is insufficient information** and recommend contacting BDO’s **compliance, legal, or cybersecurity teams**.

Your primary goal is to enable BDO Singapore teams to perform effective, regulation-aligned cybersecurity audits that reduce risk for the firm and its clients.
"""

initial_openai_message = {
    "role": "assistant",
    "content": "Hey there! 👋 I’m BDO CyberAuditBot, ready to support you with anything related to cybersecurity audits. What can I do for you today?"
}

session_docs = []


def generate_report(history):
    context = "\n\n".join(
        [f"{msg['role'].capitalize()}: {msg['content']}" for msg in history if 'content' in msg])
    prompt = f"""
    Based on the following audit context, generate a list of vulnerabilities in JSON format with the following fields for each vulnerability: 
    id, vulnerability_name, tool, timestamp, source_file_path, hostname, port, service, protocol, asset_type, business_context, description, risk_level, cvss_score, cvss_vector, cwe_id, impact, evidence, recommendation, references (as a list), document_id, document_type, chunk_text, chunk_metadata, remediation_status, ticket_id, rescan_required (boolean).
    If there is insufficient information, use default values such as 'N/A', 'Not available', empty lists for references, or false for rescan_required where applicable.

    === CONTEXT ===
    {context}
    === END CONTEXT ===

    Return only the JSON with the list of vulnerabilities.
    """

    messages = [
        {"role": "system", "content": "You are a professional cybersecurity analyst generating structured vulnerability data."},
        {"role": "user", "content": prompt}
    ]

    try:
        response_obj = client.chat.completions.create(
            model=MODEL, messages=messages)
        vulnerabilities_json = response_obj.choices[0].message.content
        vulnerabilities = json.loads(vulnerabilities_json)
    except Exception as e:
        vulnerabilities = [
            {
                "id": "VULN-000",
                "vulnerability_name": "Insufficient information",
                "tool": "N/A",
                "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "source_file_path": "N/A",
                "hostname": "N/A",
                "port": "N/A",
                "service": "N/A",
                "protocol": "N/A",
                "asset_type": "N/A",
                "business_context": "Not available",
                "description": f"Could not generate the report due to an error: {str(e)}",
                "risk_level": "N/A",
                "cvss_score": "N/A",
                "cvss_vector": "N/A",
                "cwe_id": "N/A",
                "impact": "Not available",
                "evidence": "Not available",
                "recommendation": "Review the provided history and context.",
                "references": [],
                "document_id": "N/A",
                "document_type": "N/A",
                "chunk_text": "Not available",
                "chunk_metadata": "N/A",
                "remediation_status": "Not started",
                "ticket_id": "N/A",
                "rescan_required": False
            }
        ]

    # Generate the report in Word format
    output_path = f"CyberAudit_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc = Document()

    # Cover Page
    doc.add_heading("Technical Pentesting Report", 0)
    title = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        "Generated automatically by BDO CyberAuditBot").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "1. Introduction\n2. Findings\n3. General Recommendations")
    doc.add_page_break()

    # Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report details the findings identified during an offensive security assessment (pentesting). "
        "Each vulnerability is presented with technical information, its criticality, its classification according to CVSS/CWE, "
        "and mitigation recommendations based on standards such as OWASP, MITRE ATT&CK, and Singapore regulations (PDPA, MAS)."
    )
    doc.add_page_break()

    # Findings
    doc.add_heading("2. Findings", level=1)
    for idx, vuln in enumerate(vulnerabilities, 1):
        doc.add_heading(
            f"2.{idx} {vuln.get('vulnerability_name', 'Unnamed Vulnerability')}", level=2)
        doc.add_paragraph(f"{'─' * 60}")
        doc.add_heading("🔐 VULNERABILITY DETECTED", level=3)
        doc.add_paragraph(f"{'─' * 60}")

        # General Information
        doc.add_paragraph(f"**Vulnerability ID**: {vuln.get('id', 'N/A')}")
        doc.add_paragraph(
            f"**Name**: {vuln.get('vulnerability_name', 'Not available')}")
        doc.add_paragraph(f"**Detection Tool**: {vuln.get('tool', 'N/A')}")
        doc.add_paragraph(
            f"**Detection Date**: {vuln.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
        doc.add_paragraph(
            f"**Source File**: {vuln.get('source_file_path', 'N/A')}")
        doc.add_paragraph(f"**Affected Host**: {vuln.get('hostname', 'N/A')}")
        doc.add_paragraph(
            f"**Port / Service**: {vuln.get('port', 'N/A')} / {vuln.get('service', 'N/A')}")
        doc.add_paragraph(f"**Protocol**: {vuln.get('protocol', 'N/A')}")
        doc.add_paragraph(f"**Asset Type**: {vuln.get('asset_type', 'N/A')}")
        doc.add_paragraph(
            f"**Business Context**: {vuln.get('business_context', 'Not available')}")
        doc.add_paragraph("")

        # Technical Description
        doc.add_heading("🔎 Technical Description", level=3)
        doc.add_paragraph(vuln.get('description', 'Not available'))

        # Risk Classification
        doc.add_heading("📊 Risk Classification", level=3)
        doc.add_paragraph(f"- **Risk Level**: {vuln.get('risk_level', 'N/A')}")
        doc.add_paragraph(f"- **CVSS Score**: {vuln.get('cvss_score', 'N/A')}")
        doc.add_paragraph(
            f"- **CVSS Vector**: {vuln.get('cvss_vector', 'N/A')}")
        doc.add_paragraph(f"- **CWE**: {vuln.get('cwe_id', 'N/A')}")
        doc.add_paragraph("")

        # Potential Impact
        doc.add_heading("🧠 Potential Impact", level=3)
        doc.add_paragraph(vuln.get('impact', 'Not available'))

        # Evidence
        doc.add_heading("🧪 Evidence", level=3)
        doc.add_paragraph(vuln.get('evidence', 'Not available'))
        doc.add_paragraph("")

        # Mitigation Recommendation
        doc.add_heading("✅ Mitigation Recommendation", level=3)
        doc.add_paragraph(vuln.get('recommendation', 'Not available'))
        doc.add_paragraph("")

        # References
        doc.add_heading("📚 References", level=3)
        references = vuln.get('references', [])
        if references:
            for ref in references:
                doc.add_paragraph(f"- {ref}")
        else:
            doc.add_paragraph("N/A")
        doc.add_paragraph("")

        # Internal Documentation Information
        doc.add_heading("📎 Internal Documentation Information", level=3)
        doc.add_paragraph(
            f"- **Source Document ID**: {vuln.get('document_id', 'N/A')}")
        doc.add_paragraph(
            f"- **Document Type**: {vuln.get('document_type', 'N/A')}")
        doc.add_paragraph(
            f"- **Excerpt**:\n  > {vuln.get('chunk_text', 'Not available')}")
        doc.add_paragraph(
            f"- **Chunk Metadata**: {vuln.get('chunk_metadata', 'N/A')}")
        doc.add_paragraph(
            f"- **Remediation Status**: {vuln.get('remediation_status', 'Not started')}")
        doc.add_paragraph(
            f"- **Ticket ID (optional)**: {vuln.get('ticket_id', 'N/A')}")
        doc.add_paragraph(
            f"- **Requires Rescan?**: {'Yes' if vuln.get('rescan_required', False) else 'No'}")
        doc.add_paragraph(f"{'─' * 60}")

    # General Recommendations
    doc.add_heading("3. General Recommendations", level=1)
    doc.add_paragraph(
        "It is recommended to implement the specific mitigations detailed for each vulnerability. Additionally, it is suggested to establish a continuous security assessment program, apply patches regularly, and perform rescans to verify the correction of the identified findings."
    )

    # Save the file
    doc.save(output_path)
    return output_path


def chat(message, history):
    if history is None or len(history) == 0:
        history = [initial_openai_message]

    history.append({"role": "user", "content": message})

    context_docs = vectorstore.similarity_search(message, k=10)
    context = "\n\n".join([doc.page_content for doc in context_docs])
    session_context = "\n\n".join(session_docs)

    rag_prompt = f"""
Respond exclusively based on the following context from indexed documents and recent uploaded files:

==================== BASE CONTEXT ====================
{context}
======================================================

==================== RECENT UPLOADED FILES ====================
{session_context}
==============================================================

Detect potential anomalies, vulnerabilities, or issues in the user's question and the provided context.
Classify detected vulnerabilities by CVSS and CWE if applicable.
Provide professional recommendations for mitigation.
Cite sources clearly.

User question: {message}
"""

    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": rag_prompt}
    ]

    try:
        response_obj = client.chat.completions.create(
            model=MODEL, messages=messages)
        response = response_obj.choices[0].message.content
    except Exception as e:
        response = f"Error: {str(e)}"

    history.append({"role": "assistant", "content": response})

    def history_to_tuples(history):
        tuples = []
        filtered = [m for m in history if m["role"] in ["user", "assistant"]]
        for i in range(0, len(filtered), 2):
            user_msg = filtered[i]["content"] if i < len(filtered) else ""
            assistant_msg = filtered[i+1]["content"] if i + \
                1 < len(filtered) else ""
            tuples.append((user_msg, assistant_msg))
        return tuples

    return history_to_tuples(history), history


def add_pdf_and_update_index(pdf_path):
    try:
        loader = PyPDFLoader(pdf_path)
        new_docs = loader.load()
        new_chunks = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=10).split_documents(new_docs)
        vectorstore.add_documents(new_chunks)
        for d in new_docs:
            session_docs.append(d.page_content)
        return "✅ PDF file loaded and added to the index."
    except Exception as e:
        return f"❌ Error loading the PDF file: {str(e)}"


def add_text_file_and_update_index(file_path):
    try:
        loader = TextLoader(file_path)
        new_docs = loader.load()
        new_chunks = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=10).split_documents(new_docs)
        vectorstore.add_documents(new_chunks)
        for d in new_docs:
            session_docs.append(d.page_content)
        return "✅ File loaded and added to the index."
    except Exception as e:
        return f"❌ Error loading the file: {str(e)}"


# For Gradio, use only the filename if it is in the same directory
agent_avatar = "user_avatar2.png"
user_avatar = "BDOagent2.png"
banner = "banner.png"

# Optional: warning if any image is missing
for img in [agent_avatar, user_avatar, banner]:
    abs_path = os.path.join(os.path.dirname(__file__), img)
    if not os.path.exists(abs_path):
        print(f"❌ Image not found: {abs_path}")
    else:
        print(f"✅ Image found: {abs_path}")


def start_loading():
    return gr.update(elem_classes="loading-btn loading")


def stop_loading():
    return gr.update(elem_classes="loading-btn")


def pdf_upload_click(file_obj):
    yield start_loading(), ""
    status = add_pdf_and_update_index(file_obj)
    yield stop_loading(), status


def json_upload_click(file_obj):
    yield start_loading(), ""
    status = add_text_file_and_update_index(file_obj)
    yield stop_loading(), status


def xml_upload_click(file_obj):
    yield start_loading(), ""
    status = add_text_file_and_update_index(file_obj)
    yield stop_loading(), status


def config_upload_click(file_obj):
    yield start_loading(), ""
    status = add_text_file_and_update_index(file_obj)
    yield stop_loading(), status


with gr.Blocks(css="""
    body {
        background-color: #0f172a !important;
        color: #e0e7ff;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .gr-block {
        background-color: #1e293b !important;
        border-radius: 16px;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.7);
        padding: 28px;
        margin: 20px auto;
        transition: all 0.3s ease-in-out;
        max-width: 900px;
        border: 1px solid #3b82f6;
    }
    .gr-block:hover {
        box-shadow: 0 6px 30px #3b82f6;
        transform: translateY(-3px);
    }
    h1, h2, h3, h4 {
        color: #e0e7ff;
        font-weight: 700;
    }
    .header-section {
        text-align: center;
        margin-bottom: 20px;
    }
    .header-section p {
        color: #a5b4fc;
        margin: 4px 0;
    }
    .upload-section {
        display: flex;
        gap: 16px;
        justify-content: center;
        margin-bottom: 20px;
        flex-wrap: wrap;
    }
    .upload-section > * {
        flex: 1 1 180px;
    }
    .chat-section {
        margin-bottom: 20px;
    }
    .buttons-row {
        display: flex;
        gap: 16px;
        justify-content: center;
        margin-top: 12px;
    }
    .gr-button {
        background-color: #3b82f6 !important;
        border: none !important;
        color: white !important;
        font-weight: 600 !important;
        transition: background-color 0.3s ease;
    }
    .gr-button:hover {
        background-color: #2563eb !important;
        color: #bfdbfe !important;
    }
    .gr-textbox textarea {
        background-color: #1e293b !important;
        color: #e0e7ff !important;
        border-radius: 8px !important;
        border: 1px solid #3b82f6 !important;
    }
    .gr-file input[type="file"] {
        background-color: #1e293b !important;
        color: #e0e7ff !important;
        border-radius: 8px !important;
        border: 1px solid #3b82f6 !important;
    }

    /* Loading button animation */
    .loading-btn {
        position: relative;
        overflow: hidden;
    }
    .loading-btn.loading::after {
        content: "";
        position: absolute;
        left: -50%;
        top: 0;
        width: 200%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
        animation: loadingShine 1.5s infinite;
    }

    @keyframes loadingShine {
        0% {
            left: -50%;
        }
        100% {
            left: 100%;
        }
    }
""") as demo:
    gr.Image(value=banner, width=420, height=230)
    with gr.Column(elem_classes="header-section"):
        gr.Markdown("<h1>🔐 BDO CyberAuditBot</h1>")
        gr.Markdown(
            "<p><strong>Precision Audits. Trusted Compliance.</strong></p>")
        gr.Markdown("<p>🛡️ Precision Powered by LLM & RAG</p>")
        gr.Markdown("<p>👤 <em>Designed by Carlos Egana</em> | 📍 Singapore</p>")
    gr.Markdown("---")

    with gr.Column(elem_classes="chat-section"):
        chatbot = gr.Chatbot(value=[(initial_openai_message["content"], "")],
                             height=420, type="tuples", avatar_images=(user_avatar, agent_avatar))
        history_state = gr.State([initial_openai_message])

        msg = gr.Textbox(label="Question:",
                         placeholder="Type your query here and press Enter...")
        send_btn = gr.Button("Send")

        with gr.Row(elem_classes="buttons-row"):
            with gr.Column(scale=1, min_width=150):
                report_btn = gr.Button("Generate & Download Report")
                download_report = gr.File(label="Download Report", file_types=[
                                          '.docx'], visible=False)

    with gr.Row(elem_classes="upload-section"):
        pdf_file = gr.File(label="Upload PDF",
                           type="filepath", file_types=[".pdf"])
        pdf_upload_btn = gr.Button("Load PDF", elem_classes="loading-btn")
        json_file = gr.File(label="Upload JSON",
                            type="filepath", file_types=[".json"])
        json_upload_btn = gr.Button("Load JSON", elem_classes="loading-btn")
        xml_file = gr.File(label="Upload XML",
                           type="filepath", file_types=[".xml"])
        xml_upload_btn = gr.Button("Load XML", elem_classes="loading-btn")
        config_file = gr.File(label="Upload Config (.config, .txt)",
                              type="filepath", file_types=[".config", ".txt"])
        config_upload_btn = gr.Button(
            "Load Config", elem_classes="loading-btn")

    upload_status = gr.Textbox(label="Loading status", interactive=False)

    send_btn.click(chat, inputs=[msg, history_state], outputs=[
                   chatbot, history_state]).then(lambda: "", None, msg)
    msg.submit(chat, inputs=[msg, history_state], outputs=[
               chatbot, history_state]).then(lambda: "", None, msg)

    pdf_upload_btn.click(pdf_upload_click, inputs=pdf_file,
                         outputs=[pdf_upload_btn, upload_status])
    json_upload_btn.click(json_upload_click, inputs=json_file, outputs=[
                          json_upload_btn, upload_status])
    xml_upload_btn.click(xml_upload_click, inputs=xml_file,
                         outputs=[xml_upload_btn, upload_status])
    config_upload_btn.click(config_upload_click, inputs=config_file, outputs=[
                            config_upload_btn, upload_status])

    def report_click(history):
        filepath = generate_report(history)
        return filepath, gr.update(visible=True)

    report_btn.click(report_click, inputs=history_state,
                     outputs=[download_report, download_report])

demo.launch(share=False)
